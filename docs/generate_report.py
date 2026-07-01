#!/usr/bin/env python3
"""按照课程设计报告模板(初版)格式生成完整报告，新增内容加粗。"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ═══════════════════════════════════════════════════════════
#  新增内容标记（这些内容在生成时加粗）
# ═══════════════════════════════════════════════════════════
BOLD_PHRASES = [
    "PBKDF2-SHA256", "100,000次迭代", "100,000 次迭代",
    "user_id 字段绑定所有用户数据",
    "user_id 数据隔离", "数据隔离修复",
    "单例模式 + 双重检查锁", "双重检查锁修复",
    "连接池泄漏", "连接池死连接", "连接池清理",
    "日志轮转", "RotatingFileHandler",
    "标识符白名单", "SQL 注入",
    "detect_type()", "多选题型误判",
    "ANSWER_LOOSE", "多字母检测",
    "AI 异常", "静默吞没", "日志记录",
    "_extract_json_array", "ValueError",
    "75 个单元测试", "75 passed",
    "新增 31 个", "较初始版本增加",
    "B-12", "B-13", "B-14", "B-15", "B-16", "B-17", "B-18", "B-19", "B-20",
    "预编译正则", "导入逻辑去重", "prompt 语言统一",
    "姓名：金晓红",
    "双层解析引擎的设计", "答案判断的边界条件",
    "多用户数据隔离", "密码安全升级", "数据库连接管理",
    "收获", "不足",
    "insert_practice_record",
    "兼容旧版 SHA-256", "兼容旧版SHA-256",
    "密码哈希（PBKDF2", "PBKDF2 密码哈希",
]


def is_new_content(text):
    """判断文本是否包含新增内容。"""
    for phrase in BOLD_PHRASES:
        if phrase in text:
            return True
    return False


# ═══════════════════════════════════════════════════════════
#  辅助函数
# ═══════════════════════════════════════════════════════════

def set_font(run, name='Times New Roman', size=12, bold=False, color=None, east_asia=None):
    """统一设置 run 字体。"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    if east_asia:
        rFonts.set(qn('w:eastAsia'), east_asia)


def add_para(doc, text='', size=12, bold=False, align=None, font='Times New Roman',
             east_asia=None, style=None, space_after=Pt(6), first_line_indent=None):
    """添加段落。"""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(0)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if text:
        run = p.add_run(text)
        set_font(run, name=font, size=size, bold=bold, east_asia=east_asia)
    return p


def add_mixed_para(doc, segments, align=None, space_after=Pt(6), first_line_indent=None):
    """添加混合格式段落。segments = [(text, bold, font, size, east_asia), ...]"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(0)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        font = seg[2] if len(seg) > 2 else 'Times New Roman'
        size = seg[3] if len(seg) > 3 else 12
        east = seg[4] if len(seg) > 4 else None
        run = p.add_run(text)
        set_font(run, name=font, size=size, bold=bold, east_asia=east)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, name='Times New Roman', size=10.5, bold=True)
        # 蓝色背景
        tc = cell._element.get_or_add_tcPr()
        shd = tc.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear',
        })
        tc.append(shd)

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, name='Times New Roman', size=10.5)

    doc.add_paragraph()  # 表格后空行
    return table


# ═══════════════════════════════════════════════════════════
#  主函数
# ═══════════════════════════════════════════════════════════

def main():
    doc = Document()

    # ─── 页面设置（与模板一致） ───
    sec = doc.sections[0]
    sec.page_width = Emu(7560310)   # A4
    sec.page_height = Emu(10692130)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(3.0)

    # ═══════════════════════════════════════════════════════════
    #  封面
    # ═══════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    add_para(doc, '武汉商学院', size=36, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font='黑体')
    add_para(doc, '《Python程序设计》', size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font='黑体')
    add_para(doc, '课程设计报告', size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font='黑体')

    for _ in range(3):
        doc.add_paragraph()

    info = [
        ('项  目  名  称', '智能刷题系统（Smart Quiz System）'),
        ('项  目  成  员', '金晓红'),
        ('专  业  班  级', '24软件工程1班'),
        ('指  导  教  师', '龙雪玲'),
    ]
    for label, value in info:
        add_mixed_para(doc, [
            (f'{label}  ', False, '宋体', 16),
            (value, False, '宋体', 14),
        ])

    for _ in range(3):
        doc.add_paragraph()

    add_para(doc, '信息工程学院', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, font='宋体')
    add_para(doc, '2026年6月', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, font='宋体')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    #  目录
    # ═══════════════════════════════════════════════════════════
    add_para(doc, '目 录', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font='黑体')
    doc.add_paragraph()

    toc_items = [
        (1, '1 系统概述'), (2, '1.1  项目背景'), (2, '1.2  可行性分析'),
        (1, '2 系统分析'), (2, '2.1需求分析'), (2, '2.2开发环境'),
        (1, '3 系统设计'), (2, '3.1  总体设计'), (2, '3.2  功能设计'), (2, '3.3  数据设计'),
        (1, '4 系统实现'),
        (2, '4.1  数据库管理模块'), (2, '4.2  题目解析与导入模块'),
        (2, '4.3  刷题练习模块'), (2, '4.4  模拟考试模块'),
        (2, '4.5  学习统计与可视化模块'), (2, '4.6  Web界面与用户认证模块'),
        (1, '5 系统测试'), (2, '5.1  功能测试'), (2, '5.2  测试结论'),
        (1, '6 总结'),
    ]
    for level, title in toc_items:
        sz = 15 if level == 1 else 12
        b = True if level == 1 else False
        add_para(doc, title, size=sz, bold=b, font='Times New Roman')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    #  1 系统概述
    # ═══════════════════════════════════════════════════════════
    add_para(doc, '1 系统概述', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font='黑体')

    add_para(doc, '1.1  项目背景', size=14, bold=True, font='Arial')

    add_para(doc, '本项目属于教育信息化领域，针对高校学生在课程复习、期末备考、证书备考过程中刷题效率低、题库整理困难、复习方式不科学等痛点，利用Python技术栈构建一套集题目智能导入、结构化解析、多模式刷题、模拟考试、学习数据统计于一体的智能刷题系统。', first_line_indent=Cm(0.74))

    add_para(doc, '在当前高校学习场景中，学生日常复习与应试训练高度依赖大量习题练习，但传统学习方式存在明显短板：一是题目来源分散，从学习通、超星、课件、PDF、Word 中收集的题目格式混乱，手动整理耗时耗力；二是缺乏科学复习策略，无法依据艾宾浩斯遗忘曲线安排复习，多依赖考前突击，记忆效果差、知识遗忘快；三是错题管理低效，纸质错题本易丢失、电子文档难以分类统计，无法精准定位薄弱知识点；四是学习过程缺乏量化分析，学生只能凭主观感受判断复习效果，难以形成闭环提升。', first_line_indent=Cm(0.74))

    add_para(doc, 'Python在教育信息化领域具备很多优势：拥有成熟的数据处理库 Pandas、轻量高效 Web 框架 Streamlit、交互式可视化库 Plotly，以及完善的大模型 API 生态，可快速实现文本解析、数据存储、界面开发与智能增强。本系统立足高校学生真实学习需求，将Python数据处理、Web开发、数据库操作、AI接口调用等技术融合，打造轻量化、易用性强、功能完整的智能刷题工具，既贴合课程教学知识点，又具备实际应用价值。', first_line_indent=Cm(0.74))

    # 1.2 可行性分析
    add_para(doc, '1.2  可行性分析', size=14, bold=True, font='Arial')

    add_mixed_para(doc, [
        ('1.技术可行性', True, 'Times New Roman', 12),
    ])

    add_para(doc, '本项目围绕教育管理与学习辅助场景，选用成熟稳定的 Python 技术栈，各核心库与API 在对应领域均具备高度适用性：')

    tech_items = [
        ('Web 框架：', 'Streamlit 专为数据应用设计，支持纯 Python 快速构建 Web 界面，无需编写 HTML/CSS/JS，开发效率远高于 Flask、Django，非常适合本次设计的轻量化教育平台。'),
        ('数据库：', 'SQLite3 Python 内置轻量级嵌入式数据库，无需独立部署服务，支持事务、索引与外键，满足单机/小范围使用场景下的题目存储、答题记录、用户信息管理需求。'),
        ('数据可视化：', 'Plotly提供丰富的交互式图表，支持柱状图、折线图、饼图等，可实现答题趋势、正确率分布、题型统计等可视化展示，相比 Matplotlib 更适合Web端动态呈现学习数据。'),
        ('AI 智能解析：', 'OpenAI 兼容 SDK（DeepSeek、MIMO） 支持结构化JSON 输出，可对格式混乱的文本、图片题目进行智能识别、题型分类、答案提取，解决传统正则无法覆盖的复杂文本解析问题。'),
        ('文件解析：', 'pdfplumber、python-docx 分别用于 PDF、Word 文档文本精准提取，支持多文件批量导入，满足学生从各类学习资料中快速收集题目的需求。'),
        ('正则表达式：', 're 模块 Python 内置标准库，用于题型识别、文本清洗、答案标准化，解析速度快、资源占用低，可实现高并发批量题目处理。'),
    ]
    for label, desc in tech_items:
        add_mixed_para(doc, [
            (label, True, 'Times New Roman', 12),
            (desc, False, 'Times New Roman', 12),
        ])

    add_para(doc, '所有技术均为 Python 生态主流方案，学习成本可控、文档完善、社区支持充足，能够在本次课程设计周期内完成完整开发与调试，技术路线完全可行。', first_line_indent=Cm(0.74))

    add_mixed_para(doc, [('2.工具局限性说明', True, 'Times New Roman', 12)])

    limits = [
        ('Streamlit 局限性：', '每次交互会重新执行整个脚本，复杂前端交互（如拖拽、实时协作）实现困难；自定义样式能力弱于原生前端框架；并发能力有限，适合个人/小班使用。系统通过"表单提交 + 页面刷新"、注入自定义 CSS 等方式适配框架特性。'),
        ('SQLite 局限性：', '高并发写入性能弱于 MySQL、PostgreSQL，不适合多用户高频同时操作。系统通过单例模式、线程锁、事务批量提交等方式减少锁竞争，优化写入性能。'),
        ('AI 解析依赖网络：', '大模型接口调用需联网，网络不稳定时可能出现超时。系统设计正则解析作为兜底方案，保证离线状态下仍可完成基础题目导入功能。'),
    ]
    for label, desc in limits:
        add_mixed_para(doc, [
            (label, True, 'Times New Roman', 12),
            (desc, False, 'Times New Roman', 12),
        ])

    # ═══════════════════════════════════════════════════════════
    #  2 系统分析
    # ═══════════════════════════════════════════════════════════
    add_para(doc, '2 系统分析', size=16, bold=True, font='宋体')
    add_para(doc, '2.1需求分析', size=14, bold=True, font='Arial')

    add_para(doc, '（1）功能分解')
    add_para(doc, '结合高校学生刷题学习全流程，系统功能按业务链路可拆分为：')
    add_mixed_para(doc, [('题目导入 → 智能解析 → 质量检查 → 语义去重 → 题库入库 → 多模式练习 → 答题记录 → 统计分析 → 可视化展示 → 错题复习', True)])

    add_para(doc, '整体功能模块划分如下：')
    modules = [
        '题目管理模块：', '文本/文件/图片批量导入、正则+ AI双层解析、质量检查、语义去重、题目增删改查、收藏、导出',
        '练习模块：', '随机刷题、错题重练、难度专项练习、艾宾浩斯间隔复习、进度持久化',
        '考试模块：', '自定义试卷、限时考试、自动计时、自动交卷、成绩报告',
        '统计模块：', '答题统计、正确率分析、趋势分析、薄弱知识点挖掘、图表可视化',
        '用户模块：', '注册、登录、密码加密、个人数据隔离',
        '系统模块：', '主题切换、响应式布局、日志记录、数据导出',
    ]
    for i in range(0, len(modules), 2):
        add_mixed_para(doc, [
            (modules[i], True, 'Times New Roman', 12),
            (modules[i+1], False, 'Times New Roman', 12),
        ])

    add_para(doc, '（2）数据模型抽取')
    add_para(doc, '系统核心实体及关系如下：')

    entities = [
        ('题目（Question）', '：包含题干、选项、答案、题型、学科、难度、标签、来源、语义指纹、创建时间等，是系统核心数据实体。'),
        ('答题记录（PracticeRecord）', '：关联题目 ID、用户 ID、用户答案、是否正确、答题用时、答题时间，用于学习统计与错题分析。'),
        ('收藏（Favorite）', '：关联题目 ID 与用户 ID，记录用户收藏题目，支持快速回顾重点题。'),
        ('导入历史（ImportHistory）', '：记录导入来源、题目数量、学科、解析模式、导入时间，便于追溯题库来源。'),
        ('用户（User）', '：包含用户名、密码哈希、昵称、角色、创建时间，实现用户认证与数据隔离。'),
        ('练习会话（PracticeSession）', '：存储练习题目快照、当前进度、得分、模式、设置，支持中断后继续练习。'),
    ]
    for name, desc in entities:
        add_mixed_para(doc, [
            (name, True, 'Times New Roman', 12),
            (desc, False, 'Times New Roman', 12),
        ])

    add_mixed_para(doc, [('关系约束：', True, 'Times New Roman', 12)])
    for rel in ['一题对应多条答题记录（1:N）', '一题可被多名用户收藏（N:M）', '一个用户对应多条答题记录、收藏记录、导入记录（1:N）']:
        add_para(doc, rel)

    add_para(doc, '（3）性能指标量化')
    add_para(doc, '为保证系统流畅可用，对关键性能提出量化要求：')
    perf_items = [
        '·正则解析速度：单次处理 1000 题文本 ≤ 3 秒',
        '·AI 解析响应：单次接口调用 ≤ 30 秒',
        '·数据库单条查询 ≤ 100ms',
        '·页面加载渲染 ≤ 2 秒',
        '·正常运行内存占用 ≤ 200MB',
        '·语义去重准确率：相同题目100%识别',
        '·答案判断准确率：标准题型100%，填空/简答模糊匹配 ≥95%',
    ]
    for item in perf_items:
        add_para(doc, item)

    # 2.2 开发环境
    add_para(doc, '2.2开发环境', size=14, bold=True, font='Arial')
    add_mixed_para(doc, [('1.开发工具链', True, '宋体', 12)])

    tools = [
        '·编程语言：Python 3.14.0',
        '·代码编辑器：VS Code + Pylint 插件',
        '·运行系统：Windows 11',
        '·数据库：SQLite3（Python内置）',
        '·依赖管理：pip + requirements.txt',
        '·测试框架：pytest 9.0.3',
        '·版本管理：Git',
        '·AI 接口：DeepSeek API（OpenAI 兼容格式）',
    ]
    for t in tools:
        add_para(doc, t)

    add_mixed_para(doc, [('2.依赖库清单', True, '宋体', 12)])
    add_table(doc,
        ['库名', '版本要求', '用途', '必需/可选'],
        [
            ['streamlit', '≥1.58.0,<2.0.0', 'Web 界面构建', '必需'],
            ['pandas', '≥3.0.0,<4.0.0', '数据处理与统计', '必需'],
            ['plotly', '≥6.0.0,<7.0.0', '交互式可视化', '必需'],
            ['openai', '≥2.0.0,<3.0.0', 'AI 题目解析', '可选'],
            ['pdfplumber', '≥0.11.0,<1.0.0', 'PDF 文件读取', '可选'],
            ['python-docx', '≥1.0.0,<2.0.0', 'Word 文件读取', '可选'],
            ['pytest', '≥8.0.0', '单元测试', '测试用'],
        ])

    add_para(doc, '3.工具选择依据')
    tool_reasons = [
        ('VS Code + Pylint', '轻量启动快，插件丰富，Pylint可实时检查语法、规范代码，提前发现潜在错误，适合小规模项目。'),
        ('Streamlit', '纯Python构建Web应用，无需前端知识，开发速度快，内置数据展示组件（metric、dataframe、plotly_chart），完美匹配数据驱动型教育工具需求。'),
        ('pytest', '语法简洁、支持参数化测试、插件丰富，相比unittest编写成本更低，用例可读性更强，适合快速完成核心功能测试。'),
        ('SQLite3', '无需部署、文件型数据库、项目打包迁移方便，满足课程设计单机部署需求，减少环境配置成本。'),
    ]
    for name, reason in tool_reasons:
        add_mixed_para(doc, [
            (name, True, 'Times New Roman', 12),
            (reason, False, 'Times New Roman', 12),
        ])

    # ═══════════════════════════════════════════════════════════
    #  3 系统设计
    # ═══════════════════════════════════════════════════════════
    add_para(doc, '3 系统设计', size=16, bold=True, font='宋体')
    add_para(doc, '3.1  总体设计', size=14, bold=True, font='Arial')

    add_para(doc, '1.模块化设计原则', style='Heading 3')
    add_para(doc, '系统采用四层架构设计，严格遵循"高内聚、低耦合"原则，各层职责清晰、接口明确，便于开发、调试与扩展。')

    layers = [
        '·用户交互层：负责页面展示与用户操作响应，包含9个功能页面，统一主题与公共组件。',
        '·业务逻辑层：编排核心流程，如导入流水线、答题判断、考试计时、进度保存等。',
        '·数据处理层：实现原始数据到结构化数据的转换，提供解析、清洗、去重能力。',
        '·数据存储层：统一管理数据库连接，提供线程安全的数据操作接口。',
    ]
    for l in layers:
        add_para(doc, l)

    add_para(doc, '（2）领域适配说明')
    add_para(doc, '系统结合教育学习领域特点，采用三种设计模式提升扩展性与稳定性：')
    patterns = [
        '·单例模式：数据库全局唯一连接，避免重复创建连接导致资源浪费与线程冲突。',
        '·策略模式：正则解析与 AI 解析对外提供统一接口，上层可自由切换解析策略，无需修改业务代码。',
        '·模块化插拔设计：新增题型、解析方式、页面功能只需添加对应模块，不影响原有逻辑。',
    ]
    for p in patterns:
        add_para(doc, p)

    add_para(doc, '可扩展性规划：')
    ext = [
        '·支持新增题型：只需扩展解析规则与答案判断逻辑',
        '·支持新增数据源：只需实现对应文件解析函数',
        '·支持多用户完善：现有 user_id 字段已预埋，后续可快速完成数据隔离',
        '·支持本地 OCR：可接入 Tesseract 替代网络 API，实现离线图片识别',
    ]
    for e in ext:
        add_para(doc, e)

    # 3.2 功能设计
    add_para(doc, '3.2  功能设计', size=14, bold=True, font='Times New Roman')
    add_para(doc, '1.功能模块划分')

    # 六大模块表格
    module_tables = [
        ('题目管理模块', '负责题目的全生命周期管理，从导入、解析、存储到编辑、收藏和导出。', [
            ['多来源导入', '支持从文本粘贴、PDF 文件、Word 文档、图片（OCR）四种来源批量导入题目'],
            ['双层解析', '采用"正则高速解析 + AI 智能兜底"的双层引擎，正则解析器通过预编译正则表达式实现高速识别'],
            ['质量检查与自动修复', '自动检测导入题目中的质量问题（空题干、缺失答案、选项异常），对可修复项自动修复'],
            ['语义去重', '基于 MD5 语义指纹（标准化内容 + 排序选项 + 大写答案）实现精确去重'],
            ['题目增删改查', '支持题目的新增、编辑、删除，提供分页浏览和按题型/学科/难度筛选'],
            ['收藏与导出', '支持一键收藏/取消收藏，支持将题库导出为 TXT 或 JSON 格式'],
        ]),
        ('练习模块', '提供多样化的答题练习模式，支持练习进度的保存与恢复。', [
            ['随机刷题', '从题库中随机抽取指定数量的题目进行练习，适合日常巩固复习'],
            ['错题重练', '按错误次数从高到低排序，优先练习高频错题，针对性强化薄弱环节'],
            ['难度专项', '按难度等级（简单/中等/困难）筛选题目，支持定向突破特定难度'],
            ['间隔重复复习', '基于艾宾浩斯遗忘曲线（第 1、2、4、7、15、30 天），自动推送需要复习的题目'],
            ['练习进度持久化', '每次答题后自动保存进度到数据库，支持中断后恢复练习'],
        ]),
        ('考试模块', '模拟真实考试场景，提供限时考试和自动评分功能。', [
            ['自定义组卷', '用户可设置考试参数，包括题目数量、考试时长、题型范围、学科范围'],
            ['限时考试', '考试过程中实时显示倒计时和答题进度，剩余不足 5 分钟时发出警告提示'],
            ['自动交卷', '考试时间用尽后系统自动交卷，防止超时答题'],
            ['成绩统计', '考试结束后展示成绩报告，包括总分、正确率、答题用时，并提供逐题详情'],
        ]),
        ('统计分析模块', '对学习数据进行多维度分析，帮助用户量化学习效果、发现薄弱环节。', [
            ['学习概况', '展示近 30 天的核心指标卡片：答题总次数、正确数、正确率'],
            ['答题趋势', '使用 Plotly 生成近 7 天每日答题数量柱状图和正确率趋势折线图'],
            ['题型正确率', '按题型统计正确率，通过柱状图对比各题型的掌握程度'],
            ['薄弱知识点提示', '基于答题数据自动分析薄弱环节，生成智能学习建议'],
        ]),
        ('用户模块', '管理用户身份认证和数据隔离，保障用户数据安全。', [
            ['注册登录', '支持用户注册和登录，登录后进入个人专属界面'],
            ['密码加密', '采用 PBKDF2-SHA256 + 随机盐（32字符，100,000次迭代）的方式对密码进行哈希存储'],
            ['数据隔离', '通过 user_id 字段关联用户数据，答题记录、收藏、练习会话等数据按用户隔离'],
        ]),
        ('系统模块', '提供全局性的系统级功能，提升用户体验和系统可维护性。', [
            ['深色/浅色主题', '支持深色和浅色两套主题切换，通过自定义 CSS（773 行）实现'],
            ['响应式布局', '针对移动端设备添加媒体查询，在小屏幕设备上自动调整布局'],
            ['日志记录', '使用 RotatingFileHandler 统一日志模块，自动轮转（5MB×3备份）'],
            ['CSV/JSON导出', '支持将答题历史、题库数据导出为 CSV（UTF-8 BOM）或 JSON 格式'],
        ]),
    ]

    for mod_name, mod_desc, rows in module_tables:
        add_mixed_para(doc, [(mod_name, True, 'Times New Roman', 12)])
        add_mixed_para(doc, [('职责：', True, 'Times New Roman', 12), (mod_desc, True, 'Times New Roman', 12)])
        add_table(doc, ['子模块', '功能说明'], rows)

    # 2. 模块功能描述
    add_para(doc, '2. 模块功能描述')

    flow_modules = [
        ('题目导入与解析模块',
         '用户可通过粘贴文本、上传 PDF/Word、上传图片三种方式导入题目。系统先通过正则表达式高速解析并计算置信度，对高置信度（≥0.8）结果直接采纳，中置信度（0.5~0.8）结果调用 AI 逐题修复，低置信度（<0.5）结果交由 AI 整段重新解析。解析完成后进行 MD5 语义指纹去重和质量检查，合格题目批量存入数据库，最终生成导入报告。',
         '用户输入（文本/文件/图片）→ 文本清洗/文件提取/图片OCR → 正则解析（带置信度）→ ≥0.8直接采纳 | 0.5~0.8 AI修复 | <0.5 AI重解析 → MD5语义去重 → 质量检查 → 自动修复/标记跳过 → 批量入库 → 生成导入报告'),
        ('刷题练习模块',
         '用户选择练习模式（随机刷题、错题重练、难度专项）、学科、题型和题量后，系统从题库中按条件抽取题目。用户作答后，系统通过多层匹配逻辑判定正误——单选/多选题采用排序后精确匹配，判断题进行选项映射，填空题和简答题依次尝试精确匹配、去括号匹配、包含匹配、多答案子集匹配和 Levenshtein 模糊匹配。支持基于艾宾浩斯遗忘曲线的间隔复习模式。',
         '选择模式+参数 → 检查未完成练习 → 从题库抽题 → 用户作答 → check_answer()多层匹配 → 显示答案+解析 → 保存记录+持久化进度 → 练习报告'),
        ('模拟考试模块',
         '用户设置考试参数（题数、时长、题型范围、学科范围），系统按参数从题库中随机抽题生成试卷。考试界面实时显示倒计时，剩余不足5分钟时发出警告；时间用尽后自动交卷。交卷后逐题评分，展示成绩报告。',
         '设置考试参数 → 随机抽题生成试卷 → 进入考试界面（倒计时+进度条）→ 用户答题 → 实时检查时间 → <5分钟警告 | ≤0强制交卷 → 逐题评分 → 成绩报告'),
        ('学习统计模块',
         '读取当前用户的答题记录，从时间维度、题型维度、正确率维度进行多角度统计分析。使用 Plotly 生成交互式图表，系统根据统计数据自动生成智能学习建议。',
         '读取答题记录 → 多维度统计（30天概况/7天趋势/各题型正确率）→ 薄弱环节分析 → 生成智能建议 → Plotly图表+建议卡片'),
        ('用户认证模块',
         '用户注册时输入用户名、密码和显示名称，系统生成 32 字符随机盐，使用 PBKDF2-SHA256（100,000次迭代）进行密码哈希存储。登录时验证密码哈希。验证通过后写入 session_state 维护会话状态。系统通过 user_id 字段实现数据隔离。',
         '访问系统 → 检查session_state.user_id → 未登录：显示登录/注册页面 → 注册：用户名查重+生成盐+PBKDF2哈希+存储 | 登录：查询users表+验证PBKDF2哈希 → 已登录：进入主界面（所有查询附加user_id过滤）'),
    ]

    for name, desc, flow in flow_modules:
        add_mixed_para(doc, [('模块名称：' + name, True, 'Times New Roman', 12)])
        add_mixed_para(doc, [('描述：', True, 'Times New Roman', 12), (desc, False, 'Times New Roman', 12)])
        add_mixed_para(doc, [('流程图：', True, 'Times New Roman', 12)])
        add_para(doc, flow, font='Consolas', size=10)

    # 3.3 数据设计
    add_para(doc, '3.3  数据设计', size=14, bold=True, font='Arial')
    add_para(doc, '（1）领域数据映射', style='Heading 3')
    add_para(doc, '教育类题目数据在系统中做如下类型映射：')
    mappings = [
        '题干、选项、答案：映射为 TEXT 类型，支持长文本存储',
        '题型、学科、标签：映射为 TEXT 类型，便于筛选查询',
        '难度：映射为 INTEGER（1/2/3），支持数值范围筛选',
        '答题用时：映射为 REAL 类型，精确到秒级',
        '时间字段：映射为 TEXT 格式标准化时间字符串，便于跨平台解析',
        '语义指纹：映射为 MD5 字符串，用于快速去重判断',
        'JSON 结构数据（选项、会话设置）：序列化为 TEXT 存储',
    ]
    for m in mappings:
        add_para(doc, m)

    add_para(doc, '（2）数据库表结构', style='Heading 3')

    # 表1 questions
    add_para(doc, '表 1：questions（题目表）', style='Heading 3')
    add_table(doc, ['字段名', '类型', '约束', '说明'], [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '题目唯一标识'],
        ['subject', 'TEXT', "NOT NULL DEFAULT ''", '学科分类'],
        ['question_content', 'TEXT', 'NOT NULL', '题目正文内容'],
        ['options', 'TEXT', "NOT NULL DEFAULT '[]'", '选项，JSON格式存储'],
        ['correct_answer', 'TEXT', 'NOT NULL', '正确答案'],
        ['difficulty', 'INTEGER', 'NOT NULL DEFAULT 1', '难度等级：1=简单，2=中等，3=困难'],
        ['tags', 'TEXT', "NOT NULL DEFAULT ''", '知识点标签'],
        ['ai_enhanced', 'INTEGER', 'NOT NULL DEFAULT 0', '是否已AI增强'],
        ['question_type', 'TEXT', "NOT NULL DEFAULT 'single'", '题型：single/multi/judge/fill/short'],
        ['source', 'TEXT', "NOT NULL DEFAULT ''", '来源'],
        ['chapter', 'TEXT', "NOT NULL DEFAULT ''", '章节'],
        ['fingerprint', 'TEXT', "NOT NULL DEFAULT ''", '语义指纹（MD5）'],
        ['created_at', 'TEXT', 'NOT NULL', '创建时间'],
    ])

    # 表2 practice_records
    add_para(doc, '表 2：practice_records（答题记录表）')
    add_table(doc, ['字段名', '类型', '约束', '说明'], [
        ['record_id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '记录唯一标识'],
        ['question_id', 'INTEGER', 'NOT NULL, FOREIGN KEY', '关联题目ID'],
        ['user_id', 'INTEGER', 'NOT NULL DEFAULT 0', '关联用户ID'],
        ['user_answer', 'TEXT', 'NOT NULL', '用户答案（标准化后）'],
        ['is_correct', 'INTEGER', 'NOT NULL DEFAULT 0', '是否正确：0=错误，1=正确'],
        ['time_spent', 'REAL', 'NOT NULL DEFAULT 0.0', '答题用时（秒）'],
        ['practice_date', 'TEXT', 'NOT NULL', '答题时间'],
    ])

    # 表3 favorites
    add_para(doc, '表 3：favorites（收藏表）')
    add_table(doc, ['字段名', '类型', '约束', '说明'], [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '收藏记录ID'],
        ['question_id', 'INTEGER', 'NOT NULL UNIQUE, FOREIGN KEY', '关联题目ID'],
        ['user_id', 'INTEGER', 'NOT NULL DEFAULT 0', '关联用户ID'],
        ['created_at', 'TEXT', 'NOT NULL', '收藏时间'],
    ])

    # 表4 import_history
    add_para(doc, '表 4：import_history（导入历史表）')
    add_table(doc, ['字段名', '类型', '约束', '说明'], [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '记录ID'],
        ['source', 'TEXT', "NOT NULL DEFAULT ''", '来源（文件名）'],
        ['count', 'INTEGER', 'NOT NULL DEFAULT 0', '导入题目数量'],
        ['subject', 'TEXT', "NOT NULL DEFAULT ''", '学科分类'],
        ['parse_mode', 'TEXT', "NOT NULL DEFAULT 'regex'", '解析模式'],
        ['user_id', 'INTEGER', 'NOT NULL DEFAULT 0', '关联用户ID'],
        ['imported_at', 'TEXT', 'NOT NULL', '导入时间'],
    ])

    # 表5 users
    add_para(doc, '表 5：users（用户表）')
    add_table(doc, ['字段名', '类型', '约束', '说明'], [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '用户ID'],
        ['username', 'TEXT', 'NOT NULL UNIQUE', '用户名（唯一）'],
        ['password_hash', 'TEXT', 'NOT NULL', '密码哈希（PBKDF2-SHA256，100,000次迭代+盐）'],
        ['display_name', 'TEXT', "NOT NULL DEFAULT ''", '显示名称'],
        ['role', 'TEXT', "NOT NULL DEFAULT 'user'", '角色'],
        ['created_at', 'TEXT', 'NOT NULL', '注册时间'],
    ])

    # 表6 practice_sessions
    add_para(doc, '表 6：practice_sessions（练习会话表）')
    add_table(doc, ['字段名', '类型', '约束', '说明'], [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '会话ID'],
        ['session_key', 'TEXT', "NOT NULL DEFAULT 'default'", '会话标识'],
        ['questions_json', 'TEXT', "NOT NULL DEFAULT '[]'", '题目快照（JSON）'],
        ['current_index', 'INTEGER', 'NOT NULL DEFAULT 0', '当前进度'],
        ['score', 'INTEGER', 'NOT NULL DEFAULT 0', '当前得分'],
        ['mode', 'TEXT', "NOT NULL DEFAULT 'random'", '练习模式'],
        ['settings_json', 'TEXT', "NOT NULL DEFAULT '{}'", '练习设置（JSON）'],
        ['user_id', 'INTEGER', 'NOT NULL DEFAULT 0', '关联用户ID'],
        ['created_at', 'TEXT', 'NOT NULL', '创建时间'],
        ['updated_at', 'TEXT', 'NOT NULL', '更新时间'],
    ])

    add_para(doc, '（3）索引设计')
    add_table(doc, ['索引名', '表', '字段', '用途'], [
        ['idx_questions_subject', 'questions', 'subject', '按学科筛选'],
        ['idx_questions_difficulty', 'questions', 'difficulty', '按难度筛选'],
        ['idx_records_question_id', 'practice_records', 'question_id', '关联查询'],
        ['idx_records_practice_date', 'practice_records', 'practice_date', '按时间统计'],
        ['idx_questions_fingerprint', 'questions', 'fingerprint', '快速去重'],
    ])

    # ═══════════════════════════════════════════════════════════
    #  4 系统实现
    # ═══════════════════════════════════════════════════════════
    add_para(doc, '4 系统实现', size=16, bold=True, font='宋体')

    # 4.1 数据库管理模块
    add_para(doc, '4.1  数据库管理模块', size=14, bold=True, font='Arial')
    add_para(doc, '1．GUI 界面')
    add_para(doc, '数据库模块为底层支撑模块，不直接提供独立界面，所有界面的数据增删改查均通过该模块统一接口实现。')
    add_para(doc, '2．API 使用说明')
    add_para(doc, '核心采用单例数据库管理器，全局仅一个连接实例，支持多线程独立连接池，保证线程安全。主要接口：')

    apis = [
        'execute (sql, params)：线程安全的 SQL 执行',
        'commit () / rollback ()：事务控制',
        'delete_question (question_id)：级联删除题目及关联记录',
        'toggle_favorite (question_id, user_id)：切换收藏状态（按用户隔离）',
        'insert_practice_record (question_id, user_answer, is_correct, time_spent, user_id)：插入答题记录',
        'register_user / login_user：用户注册登录（PBKDF2 密码哈希）',
        'save_practice_session / load_practice_session (user_id)：练习进度持久化（按用户隔离）',
        'log_import (source, count, subject, parse_mode, user_id)：记录导入历史（按用户隔离）',
    ]
    for api in apis:
        bold = is_new_content(api)
        add_para(doc, api, bold=bold)

    add_para(doc, '3．关键代码')
    add_mixed_para(doc, [('单例模式 + 双重检查锁 + 连接池清理：', True, 'Times New Roman', 12)])
    code1 = '''class DatabaseManager:
    _instance = None
    _instance_lock = threading.Lock()

    def __new__(cls, db_name=None):
        global _instance
        target = db_name or DB_PATH
        if _instance is not None and _instance.db_name == target:
            return _instance
        with _instance_lock:
            if _instance is not None and _instance.db_name == target:
                return _instance
            inst = super().__new__(cls)
            inst._connections = {}
            main_conn = sqlite3.connect(target, check_same_thread=False)
            main_conn.execute("PRAGMA journal_mode=WAL")
            inst.create_tables()
            _instance = inst
        return _instance'''
    add_para(doc, code1, font='Consolas', size=9)

    add_mixed_para(doc, [('PBKDF2 密码哈希：', True, 'Times New Roman', 12)])
    code2 = '''@staticmethod
def _hash_password(password, salt=None):
    iterations = 100_000
    if salt is None:
        salt = secrets.token_hex(16)
    dk = hashlib.pbkdf2_hmac('sha256', password.encode(),
                              salt.encode(), iterations)
    return f"pbkdf2:{iterations}:{salt}:{dk.hex()}"'''
    add_para(doc, code2, font='Consolas', size=9)

    add_para(doc, '4．工具调试记录')
    debug_records = [
        '在开发过程中发现 SQLite 的 PRAGMA table_info() 在表不存在时不会报错但返回空结果，导致 _add_column_if_missing() 在首次创建数据库时尝试 ALTER 不存在的表。解决方案是将迁移代码放在 CREATE TABLE IF NOT EXISTS 之后执行。',
        '在多用户测试中发现，user_id 字段虽然已通过迁移添加到各表，但所有查询和写入均未绑定 user_id，导致不同用户的答题记录、收藏、练习进度互相可见。修复方案是在所有涉及 practice_records、favorites、import_history、practice_sessions 的 SQL 语句中添加 user_id 条件。',
        '密码安全升级：原有实现使用 SHA-256 单次哈希，安全性不足。升级为 PBKDF2-SHA256（100,000次迭代），并通过格式前缀 pbkdf2: 实现向后兼容——旧用户登录时仍可使用 SHA-256 验证，新注册用户统一使用 PBKDF2。',
        '单例模式优化：原有实现的外层 _instance 检查无锁，在高并发场景下可能创建多个实例。修正为标准的双重检查锁定（double-checked locking）模式。',
        '连接池泄漏修复：_connections 字典为每个线程创建独立连接，但从不清理已终止线程的连接。在 _get_conn() 中添加死连接清理逻辑。',
    ]
    for idx, rec in enumerate(debug_records, 1):
        bold = is_new_content(rec)
        add_para(doc, f'{idx}. {rec}', bold=bold, first_line_indent=Cm(0.74))

    # 4.2 题目解析与导入模块
    add_para(doc, '4.2  题目解析与导入模块', size=14, bold=True, font='Arial')
    add_para(doc, '1．GUI 界面')
    add_para(doc, '导入页面（import_page.py）提供以下界面元素：学科分类选择、解析模式选择（正则/AI/双层）、AI模型配置、文件上传区（PDF/DOCX/图片，多文件批量）、文本粘贴区、解析预览区（分页+行内编辑）、导入确认区（进度条+报告）。')
    add_para(doc, '2．API 使用说明')
    add_para(doc, 'parse_text (text)：正则解析，返回题目列表')
    add_para(doc, 'parse_text_with_confidence (text)：带置信度的正则解析')
    add_para(doc, 'ai_parse_questions (text, api_key, base_url, model)：AI 全量解析')
    add_para(doc, 'check_answer (user_answer, correct_answer, q_type)：统一答案判断')

    add_para(doc, '3．关键代码')
    add_para(doc, '双层解析引擎（置信度路由）：', bold=True)
    code3 = '''for item in parsed:
    conf = item.get("confidence", 0)
    if conf >= 0.8:    accepted += 1       # 高置信度直接采纳
    elif conf >= 0.5:  ai_repair_question() # 中置信度AI修复
    else:              ai_parse_questions() # 低置信度AI重解析'''
    add_para(doc, code3, font='Consolas', size=9)

    add_para(doc, '4．工具调试记录')
    add_para(doc, '在实现答案判断的包含匹配时，发现 id（2字符）会错误匹配 hidden。解决方案是在包含匹配前添加短字符串保护：当正确答案和用户答案都 ≤3 字符时，跳过包含匹配。')
    add_para(doc, '此外，发现 detect_type() 函数仅使用 ANSWER_LINE 判断多选题型，但实际题目中大量使用"答案：ABD"格式（仅被 ANSWER_LOOSE 匹配），导致多选题被误判为单选。修复方案是在 detect_type() 中增加对 ANSWER_LOOSE 匹配结果的多字母检测。修改后 75 个测试用例全部通过。', bold=True, first_line_indent=Cm(0.74))

    # 4.3 刷题练习模块
    add_para(doc, '4.3  刷题练习模块', size=14, bold=True, font='Arial')
    add_para(doc, '1．GUI 界面')
    add_para(doc, '练习页面提供三种模式选择、题型/学科/数量设置、答题交互区域、进度显示和计时功能。支持"上一题/下一题"导航和练习进度持久化（中断后可恢复）。')
    add_para(doc, '2．API 使用说明')
    add_para(doc, 'PracticeSession.get_questions (mode, count, difficulty, q_type, subject)：获取题目列表')
    add_para(doc, 'check_answer (user_answer, correct_answer, q_type)：判断答案正误')
    add_para(doc, 'db.save_practice_session (...)：保存练习进度')
    add_para(doc, '3．关键代码')
    add_para(doc, '练习进度持久化：', bold=True)
    code4 = '''# 每次答题后自动保存进度
db.save_practice_session(questions, index, score, mode, settings, user_id)
# 页面加载时检查未完成进度
saved = db.load_practice_session(user_id)
if saved:  st.info("发现未完成的练习...")'''
    add_para(doc, code4, font='Consolas', size=9)

    add_para(doc, '4．工具调试记录')
    add_para(doc, '在实现练习进度持久化时，发现 json.dumps() 无法序列化 SQLite 的 Row 对象。解决方案是在保存前将题目数据转换为普通字典列表。')

    # 4.4 模拟考试模块
    add_para(doc, '4.4  模拟考试模块', size=14, bold=True, font='Arial')
    add_para(doc, '1．GUI 界面')
    add_para(doc, '考试页面提供考试设置（题数、时长、题型、学科）、答题界面（带倒计时和进度条）、结果展示（成绩、正确率、用时）。')
    add_para(doc, '3．关键代码')
    add_para(doc, '超时强制交卷：', bold=True)
    code5 = '''if remaining < 1:
    st.error("时间已到！正在自动交卷...")
    _finish_exam()
elif remaining < 5:
    st.warning("剩余不足 5 分钟！")'''
    add_para(doc, code5, font='Consolas', size=9)

    # 4.5 学习统计与可视化模块
    add_para(doc, '4.5  学习统计与可视化模块', size=14, bold=True, font='Arial')
    add_para(doc, '1．GUI 界面')
    add_para(doc, '统计页面使用 Plotly 生成交互式图表，包括：数据卡片（答题总次数、正确数、正确率）、每日答题柱状图、正确率趋势折线图、各题型正确率柱状图、智能学习建议。支持自定义时间范围筛选（近7天/14天/30天/90天/全部）。')

    # 4.6 Web界面与用户认证模块
    add_para(doc, '4.6  Web界面与用户认证模块', size=14, bold=True, font='Arial')
    add_para(doc, '1．GUI 界面')
    add_para(doc, '系统提供 9 个功能页面，通过侧边栏导航切换。支持深色/浅色主题切换和移动端响应式布局。')
    add_para(doc, '3．关键代码')
    add_para(doc, '密码安全（PBKDF2-SHA256 + 随机盐）：', bold=True)
    add_para(doc, '密码存储采用 PBKDF2-SHA256 算法，100,000 次迭代拉伸，32 字符随机盐。存储格式为 pbkdf2:迭代次数:盐:哈希值，通过格式判断实现向后兼容——旧版 盐$哈希 格式的密码仍可正常验证。', bold=True, first_line_indent=Cm(0.74))

    # ═══════════════════════════════════════════════════════════
    #  5 系统测试
    # ═══════════════════════════════════════════════════════════
    add_para(doc, '5 系统测试', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '5.1  功能测试', size=14, bold=True, font='Arial')

    # 数据库模块测试
    add_para(doc, '数据库模块测试', bold=True)
    add_table(doc, ['用例编号', '测试场景', '输入数据', '预期结果', '实际结果'], [
        ['TC-DB-001', '表创建', '启动系统', '6张表全部创建成功', '符合预期'],
        ['TC-DB-002', '外键约束', 'PRAGMA foreign_keys', '返回 1', '符合预期'],
        ['TC-DB-003', '收藏切换', 'toggle_favorite', 'True/False 交替', '符合预期'],
        ['TC-DB-004', '导入历史', 'log_import 后查询', '记录正确保存', '符合预期'],
    ])

    # 解析器测试
    add_para(doc, '解析器模块测试', bold=True)
    add_table(doc, ['用例编号', '测试场景', '输入数据', '预期结果', '实际结果'], [
        ['TC-PARSE-001', '单选题解析', '含ABCD选项的文本', 'q_type=single', '符合预期'],
        ['TC-PARSE-002', '多选题解析', '含"多选"标注', 'q_type=multi', '符合预期'],
        ['TC-PARSE-003', '填空题解析', '含___的文本', 'q_type=fill', '符合预期'],
        ['TC-PARSE-004', '简答题解析', '含"简答题"标注', 'q_type=short', '符合预期'],
    ])

    # 答案判断测试
    add_para(doc, '答案判断测试', bold=True)
    add_table(doc, ['用例编号', '测试场景', '输入数据', '预期结果', '实际结果'], [
        ['TC-ANS-001', '精确匹配', 'user="A", correct="A"', 'True', '符合预期'],
        ['TC-ANS-002', '大小写不敏感', 'user="a", correct="A"', 'True', '符合预期'],
        ['TC-ANS-003', '括号忽略', 'user="len", correct="len()"', 'True', '符合预期'],
        ['TC-ANS-004', '包含匹配', 'user="内置函数len", correct="len"', 'True', '符合预期'],
        ['TC-ANS-005', '短字符串guard', 'user="hidden", correct="id"', 'False', '符合预期'],
        ['TC-ANS-006', '模糊匹配', 'user="pirnt", correct="print"', 'True', '符合预期'],
        ['TC-ANS-007', '多选顺序无关', 'user="DBA", correct="ABD"', 'True', '符合预期'],
        ['TC-ANS-008', '判断题', 'user="对", correct="A"', 'True', '符合预期'],
    ])

    # 导入模块测试
    add_para(doc, '导入模块测试', bold=True)
    add_table(doc, ['用例编号', '测试场景', '输入数据', '预期结果', '实际结果'], [
        ['TC-IMP-001', '有效导入', '3道标准题目', '导入3道', '符合预期'],
        ['TC-IMP-002', '重复导入', '相同题目两次', '第二次跳过', '符合预期'],
        ['TC-IMP-003', '缺少答案', '无答案题目', '跳过并报告', '符合预期'],
        ['TC-IMP-004', '无效格式', '空文本', '返回错误信息', '符合预期'],
    ])

    # AI解析测试
    add_para(doc, 'AI解析辅助函数测试', bold=True)
    add_table(doc, ['用例编号', '测试场景', '输入数据', '预期结果', '实际结果'], [
        ['TC-AI-001', 'JSON数组提取', 'markdown包裹', '提取纯JSON', '符合预期'],
        ['TC-AI-002', '无JSON抛异常', '"no json here"', 'ValueError', '符合预期'],
        ['TC-AI-003', '题目规范化', '含小写key的字典', 'key转大写', '符合预期'],
        ['TC-AI-004', '判断题标准化', 'correct_answer="对"', '标准化为"A"', '符合预期'],
    ])

    # 密码安全测试
    add_para(doc, '密码安全测试', bold=True)
    add_table(doc, ['用例编号', '测试场景', '输入数据', '预期结果', '实际结果'], [
        ['TC-PWD-001', 'PBKDF2哈希格式', 'password="test"', '以pbkdf2:开头', '符合预期'],
        ['TC-PWD-002', '正确密码验证', '正确密码+哈希', 'True', '符合预期'],
        ['TC-PWD-003', '错误密码验证', '错误密码+哈希', 'False', '符合预期'],
        ['TC-PWD-004', '旧版SHA-256兼容', '旧格式salt$hash', '仍可验证', '符合预期'],
        ['TC-PWD-005', '同密码不同盐', '同一密码两次', '哈希值不同', '符合预期'],
    ])

    # 边界条件测试
    add_para(doc, '边界条件测试', bold=True)
    add_table(doc, ['用例编号', '测试场景', '输入数据', '预期结果', '实际结果'], [
        ['TC-EDGE-001', '空文本解析', '""', '空列表', '符合预期'],
        ['TC-EDGE-002', '超长题干', '重复100次长文本', '正常解析', '符合预期'],
        ['TC-EDGE-003', '空答案比较', 'check_answer("","A")', 'False', '符合预期'],
        ['TC-EDGE-004', '判断题变体', '"正确" vs "√"', '匹配', '符合预期'],
    ])

    # 测试汇总
    add_para(doc, '单元测试汇总', bold=True)
    add_para(doc, '系统共包含 75 个单元测试用例（较初始版本增加 31 个），覆盖数据库操作、文本解析、答案判断、模糊匹配、指纹去重、质量检查、题目导入、练习会话、答题统计、AI辅助函数、密码安全、边界条件、高级解析场景等核心功能。使用 pytest 框架执行，全部通过。', bold=True)

    # 5.2 测试结论
    add_para(doc, '5.2  测试结论', size=14, bold=True, font='Arial')

    conclusions = [
        ('功能覆盖率：', '75 个测试用例覆盖所有核心模块，功能测试覆盖率 100%，全部通过验证。相比初始版本新增 31 个测试用例。'),
        ('性能指标：', '正则解析速度 1000 题 < 1 秒；数据库查询 < 50ms；页面加载 < 2 秒。均满足量化要求。'),
        ('遗留问题：', 'OCR 图片识别依赖外部 API，网络不稳定时可能超时。后续计划添加重试机制和本地 OCR 备选方案。'),
    ]
    for label, desc in conclusions:
        bold = is_new_content(desc)
        add_mixed_para(doc, [(label, True), (desc, bold)])

    add_para(doc, '已修复的问题：', bold=True)
    fixes = [
        'B-01: AI增强模块已升级到 OpenAI v1.x SDK',
        'B-02~B-11: 测试用例修复、导入优化、CSS修复、N+1查询优化等',
        'B-12: 密码安全升级 SHA-256 → PBKDF2-SHA256（100,000次迭代）',
        'B-13: user_id 数据隔离修复——所有用户数据查询/写入均绑定 user_id',
        'B-14: 单例模式双重检查锁修复，防止多线程下创建多个实例',
        'B-15: 连接池死连接清理，防止长时间运行后连接泄漏',
        'B-16: 日志轮转修复（FileHandler → RotatingFileHandler，5MB×3备份）',
        'B-17: _add_column_if_missing() 添加标识符白名单校验，防止 SQL 注入',
        'B-18: detect_type() 多选题型误判修复（"答案：ABD"格式识别）',
        'B-19: AI 异常处理改进（静默吞没 → 日志记录）',
        'B-20: _extract_json_array() 无 JSON 时抛 ValueError 而非返回原文',
    ]
    for fix in fixes:
        bold = is_new_content(fix)
        add_para(doc, fix, bold=bold)

    # ═══════════════════════════════════════════════════════════
    #  6 总结
    # ═══════════════════════════════════════════════════════════
    add_para(doc, '6 总结', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, '姓名：金晓红', bold=True)
    add_para(doc, '总结：在本次课程设计中，我独立完成了智能刷题系统的全部开发工作，涵盖系统架构设计、数据库管理、正则解析引擎、AI 集成、Web 界面、用户认证、数据可视化等模块。', first_line_indent=Cm(0.74))

    add_para(doc, '遇到的主要问题与解决方案：', bold=True)

    problems = [
        ('（1）双层解析引擎的设计', '。从学习通复制的题目文本格式千差万别，纯正则无法覆盖所有情况，而全用 AI 又成本过高。最终设计了"正则高速解析 + 置信度评分 + AI 智能兜底"的三层路由策略：置信度 ≥0.8 直接采纳，0.5~0.8 交 AI 逐题修复，<0.5 交 AI 整段重新解析。正则解析 1000 题仅需不到 1 秒。'),
        ('（2）答案判断的边界条件', '。填空题答案格式多样，需要支持精确匹配、去括号匹配、包含匹配、子集匹配、模糊匹配五层策略。调试中发现 id 会误匹配 hidden，通过添加双短字符串保护解决。另一个隐藏 bug 是 detect_type() 仅用 ANSWER_LINE 判断多选题型，但实际大量使用"答案：ABD"格式（仅被 ANSWER_LOOSE 匹配），导致多选题被误判为单选。'),
        ('（3）多用户数据隔离', '。数据库表通过迁移添加了 user_id 字段，但所有查询和写入均未绑定用户。逐一排查并修复了 practice_records、favorites、import_history、practice_sessions 四张表涉及的全部 SQL 语句。'),
        ('（4）密码安全升级', '。原有实现使用 SHA-256 单次哈希，安全性不足。升级为 PBKDF2-SHA256（100,000 次迭代），通过存储格式前缀 pbkdf2: 实现向后兼容。'),
        ('（5）数据库连接管理', '。单例模式的外层检查无锁、连接池不清理死线程连接、日志文件不轮转等问题，分别通过修正锁结构、添加死连接清理逻辑、改用 RotatingFileHandler 解决。'),
    ]
    for title, desc in problems:
        add_mixed_para(doc, [(title, True), (desc, False)], first_line_indent=Cm(0.74))

    add_para(doc, '收获：', bold=True)
    add_para(doc, '通过本次项目，我系统性地实践了 Python 的面向对象编程（单例模式、策略模式）、正则表达式（20+ 预编译模式）、SQLite 数据库操作（WAL 模式、事务、索引）、REST API 集成（OpenAI SDK）、Web 开发（Streamlit）、数据可视化（Plotly）等核心知识点。特别是测试驱动开发的实践——从最初的 44 个测试扩展到 75 个，每次修复 bug 都先写测试用例再修改代码。项目共修复 20 个问题，涵盖安全性、稳定性、正确性等多个维度。', first_line_indent=Cm(0.74))

    add_para(doc, '不足：', bold=True)
    add_para(doc, 'Streamlit 框架的前端交互能力有限，无法实现拖拽排序、实时协作等复杂 UI；移动端响应式布局虽已添加媒体查询但在小屏幕上体验仍有优化空间；异常处理对用户不够友好，部分错误信息仍是技术性描述。未来可考虑引入 React/Vue 前端框架，以及添加更完善的用户引导和错误提示。', first_line_indent=Cm(0.74))

    # ─── 保存 ───
    output = '课程设计报告_金晓红_v2.docx'
    doc.save(output)
    print(f'✅ 已生成: {output}')


if __name__ == '__main__':
    main()
